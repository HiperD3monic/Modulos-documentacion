"""Convierte el reporte de auditoría Markdown a Word (.docx)"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Estilos base
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

# Helper para agregar tabla desde datos
def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

# ═══ TÍTULO ═══
title = doc.add_heading('Reporte de Auditoría: Pickings Duplicados', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
info_items = [
    ('Fecha del reporte:', '1 de mayo de 2026'),
    ('Módulo afectado:', 'pos_return (Intercambios POS)'),
    ('Criticidad:', 'Media — Afecta inventario pero no facturación'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    run_label = p.add_run(label + ' ')
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_value = p.add_run(value)
    run_value.font.size = Pt(10)

# ═══ SECCIÓN 1 ═══
doc.add_heading('1. Descripción del Bug', level=1)

doc.add_heading('¿Qué ocurrió?', level=2)
doc.add_paragraph(
    'El módulo de intercambios del Punto de Venta permite que un cliente devuelva '
    'un producto y reciba uno nuevo, pagando la diferencia de precio si aplica. '
    'Al realizar este proceso, el sistema:'
)

doc.add_paragraph('Registra los movimientos de inventario del intercambio:', style='List Number')
p = doc.add_paragraph('', style='List Bullet 2')
p.add_run('Recepción: ').bold = True
p.add_run('registra la entrada del producto devuelto al almacén')
p = doc.add_paragraph('', style='List Bullet 2')
p.add_run('Entrega: ').bold = True
p.add_run('registra la salida del producto nuevo al cliente')

doc.add_paragraph(
    'Si el cliente debe pagar una diferencia, el sistema genera una orden de pago '
    'en caja para cobrar el monto restante.', style='List Number'
)
doc.add_paragraph(
    'Al procesar esta orden de pago, el sistema volvía a crear movimientos de '
    'inventario para los mismos productos, duplicando los que ya se habían '
    'registrado en el paso 1.', style='List Number'
)

doc.add_heading('El problema', level=2)
doc.add_paragraph('Los movimientos de inventario se registraron dos veces, causando:')
doc.add_paragraph('Los productos nuevos salieron del inventario 2 veces (stock descontado doble)', style='List Bullet')
doc.add_paragraph('Los productos devueltos entraron al inventario 2 veces (stock incrementado doble)', style='List Bullet')

doc.add_heading('¿Cuándo se corrigió?', level=2)
doc.add_paragraph(
    'Se implementó una corrección en el módulo que evita la creación de movimientos '
    'de inventario duplicados cuando la orden proviene de un intercambio. Las 4 órdenes '
    'afectadas fueron creadas antes de que esta corrección estuviera activa.'
)

# ═══ SECCIÓN 2 ═══
doc.add_heading('2. Alcance del Impacto', level=1)
add_table(doc,
    ['Métrica', 'Valor'],
    [
        ['Total de órdenes de intercambio con pago', '21'],
        ['Órdenes afectadas por el bug', '4'],
        ['Órdenes protegidas por el fix', '17'],
        ['Pickings duplicados detectados', '8'],
        ['Estado de los pickings duplicados', 'Todos en "done"'],
        ['Período afectado', '21 al 23 de abril de 2026'],
        ['Tiendas afectadas', 'Caja Alameda, Caja Río 2.0'],
    ]
)

# ═══ SECCIÓN 3 ═══
doc.add_heading('3. Detalle de las 4 Transacciones Afectadas', level=1)

transactions = [
    {
        'title': '3.1 — Caja Alameda - 000023 (21 abril 2026, 22:38)',
        'session': 'Caja Alameda/00489',
        'new_products': [
            ['1005-78', 'Gorra 2025 Game Cap Olivo 5950', '1'],
            ['6023-89', 'Calcetines Marca Toros / Torín', '1'],
            ['6023-95', 'Calcetines Marca Toros / Isotipo Mini', '1'],
        ],
        'ret_products': [
            ['1005-235', 'Gorra 5950 Cápsula JDE Rojo 2025', '1'],
        ],
        'originals': [
            ['Y01/OUT/00002', 'Entrega', 'INT:05', '22:37:56', '—'],
            ['Y01/IN/00003', 'Recepción', '05', '22:37:56', '4130'],
        ],
        'duplicates': [
            ['Tie/POS/00316', 'Y01/Existencias → Customers', 'Caja Alameda - 000023', '22:38:27', '4132'],
            ['Tie/POS/00317', 'Customers → Y01/Existencias', 'Caja Alameda - 000023', '22:38:27', '4133'],
        ],
    },
    {
        'title': '3.2 — Caja Alameda - 000024 (21 abril 2026, 22:45)',
        'session': 'Caja Alameda/00489',
        'new_products': [
            ['1001-23', 'Jersey Stoli 2025 Blanca Institucional TJ, Hombre', '1'],
            ['1001-28', 'Jersey Stoli 2025 Blanca Institucional TJ, Mujer', '1'],
        ],
        'ret_products': [
            ['2002-26', 'Jersey Arrieta 2026 Blanco, Hombre', '1'],
        ],
        'originals': [
            ['Y01/OUT/00003', 'Entrega', 'INT:CAMBIO DE TALLA', '22:45:02', '—'],
            ['Y01/IN/00004', 'Recepción', 'CAMBIO DE TALLA', '22:45:02', '4134'],
        ],
        'duplicates': [
            ['Tie/POS/00318', 'Y01/Existencias → Customers', 'Caja Alameda - 000024', '22:45:10', '4136'],
            ['Tie/POS/00319', 'Customers → Y01/Existencias', 'Caja Alameda - 000024', '22:45:10', '4137'],
        ],
    },
    {
        'title': '3.3 — Caja Río 2.0 - 000015 (22 abril 2026, 19:44)',
        'session': 'Caja Río 2.0/00503',
        'new_products': [
            ['2005-10', 'Gorra 2026 Game TJ Institucional 3930', '1'],
        ],
        'ret_products': [
            ['9005-168', 'Gorra Rosa TJ Blanco 940', '1'],
        ],
        'originals': [
            ['R01/OUT/00007', 'Entrega', 'INT:267-6-000004', '19:44:42', '—'],
            ['R01/IN/00009', 'Recepción', '267-6-000004', '19:44:42', '4218'],
        ],
        'duplicates': [
            ['Tie/POS/00669', 'R01/Existencias → Customers', 'Caja Río 2.0 - 000015', '19:44:48', '4220'],
            ['Tie/POS/00670', 'Customers → R01/Existencias', 'Caja Río 2.0 - 000015', '19:44:48', '4221'],
        ],
    },
    {
        'title': '3.4 — Caja Alameda - 000032 (23 abril 2026, 00:17)',
        'session': 'Caja Alameda/00505',
        'new_products': [
            ['6023-154', 'Imán Jersey Platinum', '1'],
            ['9001-79', 'Jersey Stoli Mascotas Torín, Infantil', '1'],
        ],
        'ret_products': [
            ['2002-22', 'Jersey Arrieta 2026 Rojo, Infantil', '1'],
        ],
        'originals': [
            ['Y01/OUT/00006', 'Entrega', 'INT:030', '00:17:03', '—'],
            ['Y01/IN/00008', 'Recepción', '030', '00:17:03', '4233'],
        ],
        'duplicates': [
            ['Tie/POS/00327', 'Y01/Existencias → Customers', 'Caja Alameda - 000032', '00:17:10', '4235'],
            ['Tie/POS/00328', 'Customers → Y01/Existencias', 'Caja Alameda - 000032', '00:17:10', '4236'],
        ],
    },
]

for t in transactions:
    doc.add_heading(t['title'], level=2)
    p = doc.add_paragraph()
    p.add_run('Sesión: ').bold = True
    p.add_run(t['session'])

    p = doc.add_paragraph()
    p.add_run('Productos nuevos (salida):').bold = True
    add_table(doc, ['Código', 'Producto', 'Cantidad'], t['new_products'])

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Producto devuelto (entrada):').bold = True
    add_table(doc, ['Código', 'Producto', 'Cantidad'], t['ret_products'])

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Pickings originales (correctos, NO tocar):').bold = True
    add_table(doc, ['Referencia', 'Tipo', 'Origin', 'Fecha', 'ID'], t['originals'])

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Pickings duplicados (DEBEN REVERTIRSE):')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    add_table(doc, ['Referencia', 'Dirección', 'Origin', 'Fecha', 'ID'], t['duplicates'])

# ═══ SECCIÓN 4 ═══
doc.add_heading('4. Resumen de IDs a Revertir', level=1)
add_table(doc,
    ['ID', 'Referencia', 'Transacción', 'Tipo de movimiento'],
    [
        ['4132', 'Tie/POS/00316', 'Alameda 000023', 'Entrega (productos nuevos)'],
        ['4133', 'Tie/POS/00317', 'Alameda 000023', 'Recepción (producto devuelto)'],
        ['4136', 'Tie/POS/00318', 'Alameda 000024', 'Entrega (productos nuevos)'],
        ['4137', 'Tie/POS/00319', 'Alameda 000024', 'Recepción (producto devuelto)'],
        ['4220', 'Tie/POS/00669', 'Río 2.0 000015', 'Entrega (producto nuevo)'],
        ['4221', 'Tie/POS/00670', 'Río 2.0 000015', 'Recepción (producto devuelto)'],
        ['4235', 'Tie/POS/00327', 'Alameda 000032', 'Entrega (productos nuevos)'],
        ['4236', 'Tie/POS/00328', 'Alameda 000032', 'Recepción (producto devuelto)'],
    ]
)

# ═══ SECCIÓN 5 ═══
doc.add_heading('5. Procedimiento de Corrección en Producción', level=1)

p = doc.add_paragraph()
run = p.add_run('⚠ IMPORTANTE: Antes de ejecutar estos pasos en producción, se recomienda realizar un respaldo completo de la base de datos.')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_heading('Paso 1 — Verificar cada picking duplicado', level=2)
doc.add_paragraph('Para cada picking reportado como duplicado:')
doc.add_paragraph('Abrir el picking duplicado según su tipo de movimiento:', style='List Number')
p = doc.add_paragraph('', style='List Bullet 2')
p.add_run('Si es una Entrega ').bold = True
p.add_run('(productos nuevos — dirección: Existencias → Customers): ir a ')
p.add_run('Inventario → Operaciones → Entregas').bold = True
p = doc.add_paragraph('', style='List Bullet 2')
p.add_run('Si es una Recepción ').bold = True
p.add_run('(producto devuelto — dirección: Customers → Existencias): ir a ')
p.add_run('Inventario → Operaciones → Recepciones').bold = True
doc.add_paragraph('En la barra de búsqueda, escribir la Referencia del picking duplicado (ej: "Tie/POS/00316") o filtrar por el Documento origen (ej: "Caja Alameda - 000023")', style='List Number')
doc.add_paragraph('Verificar que el Documento origen sea una orden POS de intercambio (ej: "Caja Alameda - 000023")', style='List Number')
doc.add_paragraph('Abrir en otra pestaña el picking original del exchange (ej: "Y01/OUT/00002") buscándolo de la misma forma en Entregas o Recepciones según corresponda', style='List Number')
doc.add_paragraph('Confirmar que ambos tienen los mismos productos y cantidades', style='List Number')
doc.add_paragraph('Verificar que el picking duplicado fue creado segundos después del original', style='List Number')

doc.add_heading('Paso 2 — Revertir los pickings duplicados', level=2)
doc.add_paragraph('Para cada picking duplicado confirmado:')
doc.add_paragraph('Abrir el picking en el formulario de Inventario', style='List Number')
doc.add_paragraph('Hacer clic en el botón "Devolver" (esquina superior izquierda)', style='List Number')
doc.add_paragraph('En el popup de devolución, verificar las cantidades y hacer clic en "Devolver"', style='List Number')
doc.add_paragraph('Se creará un picking de devolución en estado "Listo"', style='List Number')
doc.add_paragraph('Abrir el picking de devolución y hacer clic en "Validar"', style='List Number')
doc.add_paragraph('Confirmar la validación', style='List Number')

p = doc.add_paragraph()
run = p.add_run('Repetir el Paso 2 para cada uno de los 8 pickings duplicados.')
run.bold = True

doc.add_heading('Paso 3 — Verificar el inventario', level=2)
doc.add_paragraph('Después de revertir todos los pickings:')
doc.add_paragraph('Ir a Inventario → Reportes → Valoración de inventario', style='List Number')
doc.add_paragraph('Verificar que los productos afectados tengan cantidades coherentes', style='List Number')
doc.add_paragraph('Si es necesario, realizar un inventario físico de los productos específicos', style='List Number')

doc.add_heading('Paso 4 — Confirmar que el fix previene futuros duplicados', level=2)
doc.add_paragraph('Verificar que el módulo pos_return actualizado esté desplegado en producción', style='List Number')
doc.add_paragraph('Realizar un intercambio de prueba con pago de diferencia', style='List Number')
doc.add_paragraph('Verificar que solo se cree 1 par de pickings (entrega + recepción), no pickings adicionales', style='List Number')

# ═══ SECCIÓN 6 ═══
doc.add_heading('6. Productos Afectados por la Duplicación', level=1)
doc.add_paragraph('Los siguientes productos tienen movimientos de inventario duplicados que deben corregirse:')
add_table(doc,
    ['Código', 'Producto', 'Movimiento duplicado', 'Efecto en inventario'],
    [
        ['1005-78', 'Gorra 2025 Game Cap Olivo 5950', 'Salida extra x1', 'Stock subcontado (-1)'],
        ['6023-89', 'Calcetines Marca Toros / Torín', 'Salida extra x1', 'Stock subcontado (-1)'],
        ['6023-95', 'Calcetines Marca Toros / Isotipo Mini', 'Salida extra x1', 'Stock subcontado (-1)'],
        ['1005-235', 'Gorra 5950 Cápsula JDE Rojo 2025', 'Entrada extra x1', 'Stock sobrecontado (+1)'],
        ['1001-23', 'Jersey Stoli 2025 Blanca Inst. TJ, Hombre', 'Salida extra x1', 'Stock subcontado (-1)'],
        ['1001-28', 'Jersey Stoli 2025 Blanca Inst. TJ, Mujer', 'Salida extra x1', 'Stock subcontado (-1)'],
        ['2002-26', 'Jersey Arrieta 2026 Blanco, Hombre', 'Entrada extra x1', 'Stock sobrecontado (+1)'],
        ['2005-10', 'Gorra 2026 Game TJ Institucional 3930', 'Salida extra x1', 'Stock subcontado (-1)'],
        ['9005-168', 'Gorra Rosa TJ Blanco 940', 'Entrada extra x1', 'Stock sobrecontado (+1)'],
        ['6023-154', 'Imán Jersey Platinum', 'Salida extra x1', 'Stock subcontado (-1)'],
        ['9001-79', 'Jersey Stoli Mascotas Torín, Infantil', 'Salida extra x1', 'Stock subcontado (-1)'],
        ['2002-22', 'Jersey Arrieta 2026 Rojo, Infantil', 'Entrada extra x1', 'Stock sobrecontado (+1)'],
    ]
)

# ═══ SECCIÓN 7 — SEGUNDO BUG ═══
doc.add_page_break()
doc.add_heading('7. Segundo Bug Detectado: Recepciones Múltiples (Local 12)', level=1)

doc.add_heading('¿Qué ocurrió?', level=2)
doc.add_paragraph(
    'En la tienda Toro Dulces y Botanas L12, se detectaron recepciones de inventario '
    'duplicadas para dos tickets de punto de venta. A diferencia del bug original '
    '(secciones 1-6), el fix para los pickings nativos del POS ya estaba aplicado '
    'en el momento en que estas transacciones fueron procesadas.'
)

doc.add_heading('Diferencias con el Bug Original', level=2)
add_table(doc,
    ['Aspecto', 'Bug Original (Secciones 1-6)', 'Nuevo Bug (L12)'],
    [
        ['Tipo de picking duplicado', 'POS nativo (Tie/POS/xxxxx)', 'Recepciones de compra (L12/IN/xxxxx)'],
        ['Tipo de operación', 'Órdenes de PdV', 'Toro Botanas: Recepciones'],
        ['Dirección', 'Existencias → Customers', 'Vendors → L12/Existencias'],
        ['Momento de creación', 'Inmediato al procesar pago', 'Batch (~12:15 AM, horas después)'],
        ['Fix de duplicados nativos', 'No estaba aplicado', 'Ya estaba aplicado'],
        ['Duplicados por ticket', '1 par (entrega + recepción)', '2 a 4 recepciones idénticas'],
    ]
)

doc.add_heading('El problema', level=2)
doc.add_paragraph(
    'Al procesar devoluciones de tickets previamente vendidos, el sistema debía crear '
    'una única recepción por devolución para registrar la entrada del producto devuelto '
    'al almacén. En su lugar, creó múltiples recepciones idénticas (2 para el ticket '
    '000194 y 4 para el ticket 000182). Esto causa un incremento falso del inventario, '
    'ya que los productos devueltos se registran como recibidos múltiples veces.'
)

# ═══ SECCIÓN 8 ═══
doc.add_heading('8. Detalle de las 2 Transacciones Afectadas (L12)', level=1)

# --- 8.1 Ticket 000194 ---
doc.add_heading('8.1 — Toro Dulces y Botanas L12 - 000194 (2 mayo 2026)', level=2)

p = doc.add_paragraph()
p.add_run('Tienda: ').bold = True
p.add_run('Toro Dulces y Botanas L12 (Local 12)')

p = doc.add_paragraph()
p.add_run('Productos del ticket:').bold = True
add_table(doc, ['Código', 'Producto', 'Cantidad'], [
    ['TDB18', 'SABRITAS VARIEDAD', '1'],
    ['TDB08', 'FRITURAS', '1'],
])

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Entrega de venta original (correcta, NO tocar):').bold = True
add_table(doc, ['Referencia', 'Tipo', 'Origen', 'Fecha', 'Productos'], [
    ['TDB/POS/00266', 'Entrega (Órdenes de PdV)', 'Toro Dulces y Botanas L12 - 000194', '2 may, 10:07 p.m.', 'SABRITAS VARIEDAD (1) + FRITURAS (1)'],
])

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Recepción de devolución (correcta, NO tocar):').bold = True
add_table(doc, ['Referencia', 'Dirección', 'Origen', 'Fecha', 'Productos'], [
    ['L12/IN/00001', 'Vendors → L12/Existencias', '2684-36-000194', '3 may, 12:15 a.m.', 'SABRITAS VARIEDAD (1) + FRITURAS (1)'],
])

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Recepción duplicada (DEBE REVERTIRSE):')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
add_table(doc, ['Referencia', 'Dirección', 'Origen', 'Fecha', 'Productos'], [
    ['L12/IN/00002', 'Vendors → L12/Existencias', '2684-36-000194', '3 may, 12:16 a.m.', 'SABRITAS VARIEDAD (1) + FRITURAS (1)'],
])

# --- 8.2 Ticket 000182 ---
doc.add_heading('8.2 — Toro Dulces y Botanas L12 - 000182 (2 mayo 2026)', level=2)

p = doc.add_paragraph()
p.add_run('Tienda: ').bold = True
p.add_run('Toro Dulces y Botanas L12 (Local 12)')

p = doc.add_paragraph()
p.add_run('Productos del ticket:').bold = True
add_table(doc, ['Código', 'Producto', 'Cantidad'], [
    ['TDB08', 'FRITURAS', '1'],
])

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Entrega de venta original (correcta, NO tocar):').bold = True
add_table(doc, ['Referencia', 'Tipo', 'Origen', 'Fecha', 'Productos'], [
    ['TDB/POS/00254', 'Entrega (Órdenes de PdV)', 'Toro Dulces y Botanas L12 - 000182', '2 may, 8:24 p.m.', 'FRITURAS (1)'],
])

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Recepción de devolución (correcta, NO tocar):').bold = True
add_table(doc, ['Referencia', 'Dirección', 'Origen', 'Fecha', 'Productos'], [
    ['L12/IN/00003', 'Vendors → L12/Existencias', '2684-36-000182', '3 may, 12:17 a.m.', 'FRITURAS (1)'],
])

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Recepciones duplicadas (DEBEN REVERTIRSE):')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
add_table(doc, ['Referencia', 'Dirección', 'Origen', 'Fecha', 'Productos'], [
    ['L12/IN/00004', 'Vendors → L12/Existencias', '2684-36-000182', '3 may, 12:17 a.m.', 'FRITURAS (1)'],
    ['L12/IN/00005', 'Vendors → L12/Existencias', '2684-36-000182', '3 may, 12:17 a.m.', 'FRITURAS (1)'],
    ['L12/IN/00006', 'Vendors → L12/Existencias', '2684-36-000182', '3 may, 12:17 a.m.', 'FRITURAS (1)'],
])

# ═══ SECCIÓN 9 ═══
doc.add_heading('9. Resumen de IDs a Revertir (L12)', level=1)
add_table(doc,
    ['Referencia', 'Transacción', 'Productos', 'Efecto en inventario'],
    [
        ['L12/IN/00002', 'L12 - 000194', 'SABRITAS VARIEDAD (1) + FRITURAS (1)', 'Stock sobrecontado (+2)'],
        ['L12/IN/00004', 'L12 - 000182', 'FRITURAS (1)', 'Stock sobrecontado (+1)'],
        ['L12/IN/00005', 'L12 - 000182', 'FRITURAS (1)', 'Stock sobrecontado (+1)'],
        ['L12/IN/00006', 'L12 - 000182', 'FRITURAS (1)', 'Stock sobrecontado (+1)'],
    ]
)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Total de recepciones a revertir: 4')
run.bold = True

# ═══ SECCIÓN 10 ═══
doc.add_heading('10. Productos Afectados por las Recepciones Duplicadas (L12)', level=1)
doc.add_paragraph('Los siguientes productos tienen stock sobrecontado por las recepciones duplicadas:')
add_table(doc,
    ['Código', 'Producto', 'Recepciones duplicadas', 'Efecto total en inventario'],
    [
        ['TDB18', 'SABRITAS VARIEDAD', '1 recepción extra (de 000194)', 'Stock sobrecontado (+1)'],
        ['TDB08', 'FRITURAS', '1 extra (000194) + 3 extra (000182)', 'Stock sobrecontado (+4)'],
    ]
)

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Procedimiento de corrección: ').bold = True
p.add_run(
    'Seguir los mismos pasos descritos en la Sección 5. Para cada recepción duplicada, '
    'ir a Inventario → Operaciones → Recepciones, buscar por la referencia (ej: L12/IN/00001), '
    'hacer clic en "Devolver", verificar cantidades, y validar la devolución.'
)

# Guardar
output_path = os.path.join(os.path.dirname(__file__), 'reporte_pickings_duplicados.docx')
doc.save(output_path)
print("Archivo guardado en:", output_path)
